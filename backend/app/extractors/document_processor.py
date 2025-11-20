"""
Document processor for extracting text from PDF, DOCX, and TXT files
"""
import os
from pathlib import Path
from typing import List, Dict
import pdfplumber
from docx import Document
from app.extractors.rule_based import RuleBasedExtractor
from app.models.schemas import Entity


class DocumentProcessor:
    """Process documents and extract entities"""
    
    def __init__(self):
        """Initialize document processor"""
        self.extractor = RuleBasedExtractor()
    
    def process_document(self, file_path: str, filename: str) -> List[Entity]:
        """
        Process document and extract entities
        
        Args:
            file_path: Path to document file
            filename: Original filename
            
        Returns:
            List of extracted Entity objects
            
        Raises:
            ValueError: If file type is unsupported
        """
        file_ext = Path(file_path).suffix.lower()
        
        # Extract text based on file type
        if file_ext == '.pdf':
            text = self._extract_from_pdf(file_path)
        elif file_ext == '.docx':
            text = self._extract_from_docx(file_path)
        elif file_ext == '.txt':
            text = self._extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        # Extract entities
        entities = self.extractor.extract(text, source=filename)
        
        return entities
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF using pdfplumber
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        text_parts = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except Exception as e:
            raise ValueError(f"Error reading PDF: {str(e)}")
        
        return "\n\n".join(text_parts)
    
    def _extract_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX using python-docx
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        text_parts = []
        
        try:
            doc = Document(file_path)
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append("\t".join(row_text))
        
        except Exception as e:
            raise ValueError(f"Error reading DOCX: {str(e)}")
        
        return "\n".join(text_parts)
    
    def _extract_from_txt(self, file_path: str) -> str:
        """
        Extract text from plain text file
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            File contents
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            raise ValueError(f"Error reading TXT: {str(e)}")

